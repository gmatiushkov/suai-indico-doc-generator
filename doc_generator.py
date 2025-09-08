import json
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os

def set_page_layout(doc):
    """Настройка размера бумаги A4 и полей 1 дюйм."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def safe_get(dct, *keys):
    """Безопасное получение значения из вложенного словаря. Возвращает '' если нет."""
    for key in keys:
        if dct and key in dct:
            dct = dct[key]
        else:
            return ''
    return dct

def create_program_docx(data, output_dir):
    """Генерация документа программы конференции."""
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Программа\n{data.get("title","")} по кафедре № 43 компьютерных технологий и программной инженерии')
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph('Секция каф.43. «компьютерных технологий и программной инженерии»')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.runs[0]
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)

    leadership = data.get('leadership', {})
    p = doc.add_paragraph(f'Научный руководитель секции – {safe_get(leadership, "scientific_leader", "name")}')
    p.paragraph_format.left_indent = Cm(2.0)
    p.runs[0].font.size = Pt(12)
    aff = safe_get(leadership, "scientific_leader", "affiliation")
    if aff:
        doc.add_paragraph(aff).paragraph_format.left_indent = Cm(2.0)

    p = doc.add_paragraph(f'Зам. научного руководителя секции – {safe_get(leadership, "deputy_leader", "name")}')
    p.paragraph_format.left_indent = Cm(2.0)
    p.runs[0].font.size = Pt(12)
    aff = safe_get(leadership, "deputy_leader", "affiliation")
    if aff:
        doc.add_paragraph(aff).paragraph_format.left_indent = Cm(2.0)

    p = doc.add_paragraph(f'Секретарь – {safe_get(leadership, "secretary", "name")}')
    p.paragraph_format.left_indent = Cm(2.0)
    p.runs[0].font.size = Pt(12)
    aff = safe_get(leadership, "secretary", "affiliation")
    if aff:
        doc.add_paragraph(aff).paragraph_format.left_indent = Cm(2.0)

    doc.add_paragraph()

    for session in data.get('sessions', []):
        p = doc.add_paragraph(f'Заседание {session.get("number","")}.')
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

        p = doc.add_paragraph(f'{session.get("date","")}, {session.get("start_time","")}, {session.get("room_name","")}')
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        doc.add_paragraph()

        p = doc.add_paragraph('По решению руководителя секции порядок следования докладов может быть изменен.')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(14)

        doc.add_paragraph()

        index = 1
        non_submitted_or_rejected = []
        for contribution in session.get('contributions', []):
            if contribution.get('review_state') in ['not submitted', 'rejected']:
                non_submitted_or_rejected.append(contribution)
                continue
            speaker = contribution.get("speaker", {})
            full_name = speaker.get("full_name","")
            affiliation = speaker.get("affiliation","")
            p = doc.add_paragraph(f'{index}. {full_name}, {affiliation}')
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-0.63)
            p.paragraph_format.keep_with_next = True
            p.runs[0].font.size = Pt(14)

            p = doc.add_paragraph(contribution.get('title',''))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.runs[0].font.size = Pt(14)
            doc.add_paragraph()
            index += 1

        if non_submitted_or_rejected:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph('Следующие далее студенты своевременно не зарегистрировались для участия в конференции. По решению руководителя секции их доклад может быть перенесен на более позднюю дату.')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(14)

            doc.add_paragraph()

            for contribution in non_submitted_or_rejected:
                speaker = contribution.get("speaker", {})
                full_name = speaker.get("full_name","")
                affiliation = speaker.get("affiliation","")
                p = doc.add_paragraph(f'{index}. {full_name}, {affiliation}')
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                p.paragraph_format.keep_with_next = True
                p.runs[0].font.size = Pt(14)

                p = doc.add_paragraph(contribution.get('title',''))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.27)
                p.runs[0].font.size = Pt(14)
                doc.add_paragraph()
                index += 1

        if session != data.get('sessions', [])[-1]:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(os.path.join(output_dir, '1_Программа_к43.docx'))

# Для функций create_report_docx и create_publication_list_docx точно так же использовать safe_get для leadership
# и .get() для всех остальных данных

def create_report_docx(data, output_dir):
    """Генерация документа отчета конференции с адресом из JSON."""
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    p = doc.add_paragraph(f'Отчет о проведении {data.get("title","")}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph('Секция 43. Кафедра компьютерных технологий и программной инженерии')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph()

    leadership = data.get('leadership', {})
    conference_address = data.get('address', '')

    for session in data.get('sessions', []):
        p = doc.add_paragraph(f'Заседание {session.get("number","")}')
        p.runs[0].bold = True

        room_name = session.get("room_name", "")
        p = doc.add_paragraph(f'{session.get("date","")} г., {session.get("start_time","")}, {conference_address}, {room_name}')
        p.runs[0].font.size = Pt(10)

        sci_leader_name = safe_get(leadership, "scientific_leader", "name")
        sci_leader_aff = safe_get(leadership, "scientific_leader", "affiliation")
        sec_name = safe_get(leadership, "secretary", "name")
        sec_aff = safe_get(leadership, "secretary", "affiliation")

        if sci_leader_name or sci_leader_aff:
            p = doc.add_paragraph(f'Научный руководитель секции – {sci_leader_aff} {sci_leader_name}'.strip())
            p.runs[0].font.size = Pt(10)

        if sec_name or sec_aff:
            p = doc.add_paragraph(f'Секретарь – {sec_aff} {sec_name}'.strip())
            p.runs[0].font.size = Pt(10)

        doc.add_paragraph()

        doc.add_paragraph('Список докладов')
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(1.0)
        table.columns[1].width = Cm(9.0)
        table.columns[2].width = Cm(2.92)
        table.columns[3].width = Cm(3.0)

        headers = ['№ п/п', 'Фамилия и инициалы докладчика, название доклада', 'Статус (магистр / студент)', 'Решение']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(10)

        for i, contribution in enumerate(session.get('contributions', []), 1):
            row = table.add_row()
            cell = row.cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(i)).font.size = Pt(10)

            cell = row.cells[1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            speaker = contribution.get("speaker", {})
            p.add_run(f'{speaker.get("full_name","")}. {contribution.get("title","")}').font.size = Pt(10)

            cell = row.cells[2]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            aff = speaker.get("affiliation","")
            status = 'Магистрант' if 'Магистрант' in aff else 'Студент'
            gr_num = aff.split("гр.")[1].strip() if "гр." in aff else ""
            p.add_run(f'{status} {gr_num}').font.size = Pt(10)

            cell = row.cells[3]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('').font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        if sci_leader_name:
            p = doc.add_paragraph(f'Научный руководитель секции _________________ / {sci_leader_name}')
            p.runs[0].font.size = Pt(10)

        if session != data.get('sessions', [])[-1]:
            doc.add_paragraph()

    doc.save(os.path.join(output_dir, f'2_Отчет о проведении {data.get("title","")}.docx'))


def create_publication_list_docx(data, output_dir):
    """Генерация документа списка докладов для публикации."""
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph('Список представляемых к публикации докладов')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].italic = True

    doc.add_paragraph()

    leadership = data.get('leadership', {})
    sec_name = safe_get(leadership, "secretary", "name")
    sec_aff = safe_get(leadership, "secretary", "affiliation")
    p = doc.add_paragraph('Кафедра № 43 компьютерных технологий и программной инженерии')
    p.paragraph_format.left_indent = Cm(2.0)
    p.runs[0].font.size = Pt(12)

    if sec_name:
        p = doc.add_paragraph(sec_name)
        p.paragraph_format.left_indent = Cm(2.0)
        p.runs[0].font.size = Pt(12)

    if leadership.get("secretary", {}).get("email"):
        p = doc.add_paragraph(f'e-mail: {leadership["secretary"]["email"]}')
        p.paragraph_format.left_indent = Cm(2.0)
        p.runs[0].font.size = Pt(12)

    p = doc.add_paragraph('тел.:')
    p.paragraph_format.left_indent = Cm(2.0)
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    accepted_contributions = []
    for session in data.get('sessions', []):
        for contribution in session.get('contributions', []):
            if contribution.get('review_state') == 'accepted':
                accepted_contributions.append(contribution)

    for i, contribution in enumerate(accepted_contributions, 1):
        speaker = contribution.get("speaker", {})
        first_name_parts = speaker.get("first_name","").split()
        last_name = speaker.get("last_name","")
        initials = f'{first_name_parts[0][0]}. {first_name_parts[1][0]}.' if len(first_name_parts) > 1 else f'{first_name_parts[0][0]}.' if first_name_parts else ''
        speaker_name = f'{last_name} {initials}'.strip()
        p = doc.add_paragraph(f'{i}. {speaker_name}. {contribution.get("title","")}')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.runs[0].font.size = Pt(14)
        doc.add_paragraph()

    doc.add_paragraph()

    if sec_name:
        p = doc.add_paragraph(f'Руководитель УНИДС _________________ / {sec_name}')
        p.paragraph_format.left_indent = Cm(2.0)
        p.runs[0].font.size = Pt(12)

    doc.save(os.path.join(output_dir, '3_Список представляемых к публикации докладов.docx'))

def create_conference_docx(data, output_dir):
    """Создание всех трех DOCX документов."""
    create_program_docx(data, output_dir)
    create_report_docx(data, output_dir)
    create_publication_list_docx(data, output_dir)

